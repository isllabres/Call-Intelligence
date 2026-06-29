from __future__ import annotations

from pathlib import Path

from call_intel.config import get_transcripts_dir


def test_should_return_transcripts_dir_from_config(monkeypatch):
    monkeypatch.delenv("TRANSCRIPTS_DIR", raising=False)
    result = get_transcripts_dir()
    assert result.name == "transcripts"
    assert result.parent.name == "input"
    assert result.parent.parent.name == "data"


def test_should_parse_vtt_file_into_transcript_model(tmp_path):
    from call_intel.transcript_parser import parse_transcript_file

    vtt_content = """\
WEBVTT

00:00:01.000 --> 00:00:04.000
Hello, welcome to the meeting.

00:00:05.500 --> 00:00:09.200
Thanks for joining, let's get started.

00:00:10.000 --> 00:00:15.000
First topic is the project timeline.
"""
    vtt_file = tmp_path / "test 01-01-2026.vtt"
    vtt_file.write_text(vtt_content)

    transcript = parse_transcript_file(vtt_file)

    assert len(transcript.segments) == 3
    assert transcript.segments[0].start == 1.0
    assert transcript.segments[0].end == 4.0
    assert transcript.segments[0].text.strip() == "Hello, welcome to the meeting."
    assert transcript.segments[1].start == 5.5
    assert transcript.segments[1].end == 9.2
    assert transcript.segments[2].start == 10.0
    assert transcript.segments[2].end == 15.0


def test_should_raise_error_for_malformed_vtt(tmp_path):
    import pytest

    from call_intel.transcript_parser import parse_transcript_file

    bad_vtt = tmp_path / "bad 01-01-2026.vtt"
    bad_vtt.write_text("This is not a VTT file\nJust some random text\n")

    with pytest.raises(ValueError, match="missing WEBVTT header"):
        parse_transcript_file(bad_vtt)


def test_should_parse_docx_file_into_transcript_model(tmp_path):
    from docx import Document

    from call_intel.transcript_parser import parse_transcript_file

    doc = Document()
    doc.add_paragraph("Welcome everyone to the meeting.")
    doc.add_paragraph("Let's discuss the project status.")
    doc.add_paragraph("We need to finalize the timeline.")
    docx_file = tmp_path / "test 01-01-2026.docx"
    doc.save(str(docx_file))

    transcript = parse_transcript_file(docx_file)

    assert len(transcript.segments) == 1
    assert "Welcome everyone" in transcript.segments[0].text
    assert "project status" in transcript.segments[0].text
    assert "finalize the timeline" in transcript.segments[0].text


def test_should_raise_error_for_empty_docx(tmp_path):
    import pytest
    from docx import Document

    from call_intel.transcript_parser import parse_transcript_file

    doc = Document()
    docx_file = tmp_path / "empty 01-01-2026.docx"
    doc.save(str(docx_file))

    with pytest.raises(ValueError, match="no text content found"):
        parse_transcript_file(docx_file)


def test_should_parse_transcript_filename_for_project_and_date(tmp_path):
    from datetime import datetime

    from call_intel.filename import parse_recording_name

    vtt_file = tmp_path / "client_acme 15-03-2026.vtt"
    vtt_file.touch()

    project, date = parse_recording_name(vtt_file)

    assert project == "client_acme"
    assert date == datetime(2026, 3, 15)


def test_should_process_transcript_and_write_output(tmp_path, monkeypatch):
    from call_intel.pipeline import process_transcript

    (tmp_path / "data" / "output").mkdir(parents=True)
    monkeypatch.setattr("call_intel.config.PROJECT_ROOT", tmp_path)

    vtt_content = """\
WEBVTT

00:00:01.000 --> 00:00:04.000
Hello, welcome to the meeting.

00:00:05.500 --> 00:00:09.200
Let's discuss the roadmap.
"""
    transcripts_dir = tmp_path / "data" / "input" / "transcripts"
    transcripts_dir.mkdir(parents=True)
    vtt_file = transcripts_dir / "acme project 15-03-2026.vtt"
    vtt_file.write_text(vtt_content)

    record = process_transcript(
        transcript_path=vtt_file,
        skip_analysis=True,
        skip_google=True,
    )

    meta_path = Path(record.output_dir) / "meta.json"
    transcript_md = Path(record.output_dir) / "transcript.md"
    assert meta_path.exists()
    assert transcript_md.exists()
    assert record.source_path == str(vtt_file.resolve())
    assert record.project == "acme project"


def test_should_detect_unprocessed_transcripts(tmp_path, monkeypatch):
    from call_intel.watcher import find_unprocessed_transcripts

    monkeypatch.setattr("call_intel.config.PROJECT_ROOT", tmp_path)

    transcripts_dir = tmp_path / "data" / "input" / "transcripts"
    transcripts_dir.mkdir(parents=True)
    output_dir = tmp_path / "data" / "output"
    output_dir.mkdir(parents=True)

    vtt_file = transcripts_dir / "project_x 01-01-2026.vtt"
    vtt_file.write_text("WEBVTT\n\n00:00:01.000 --> 00:00:02.000\nHello\n")

    unprocessed = find_unprocessed_transcripts(transcripts_dir)
    assert len(unprocessed) == 1
    assert unprocessed[0].name == vtt_file.name


def test_should_ignore_unsupported_extensions_in_transcripts_folder(tmp_path, monkeypatch):
    from call_intel.watcher import find_unprocessed_transcripts

    monkeypatch.setattr("call_intel.config.PROJECT_ROOT", tmp_path)

    transcripts_dir = tmp_path / "data" / "input" / "transcripts"
    transcripts_dir.mkdir(parents=True)
    output_dir = tmp_path / "data" / "output"
    output_dir.mkdir(parents=True)

    (transcripts_dir / "notes.txt").write_text("some notes")
    (transcripts_dir / "report.pdf").write_text("fake pdf")

    unprocessed = find_unprocessed_transcripts(transcripts_dir)
    assert unprocessed == []


def test_should_watcher_trigger_on_new_vtt_file(tmp_path):
    import threading
    import time

    from watchdog.observers import Observer

    from call_intel.watcher import TranscriptHandler

    transcripts_dir = tmp_path / "transcripts"
    transcripts_dir.mkdir()

    triggered = threading.Event()
    triggered_path = []

    class TestHandler(TranscriptHandler):
        def _process_transcript(self, path):
            triggered_path.append(path)
            triggered.set()

    handler = TestHandler(skip_google=True)
    observer = Observer()
    observer.schedule(handler, str(transcripts_dir), recursive=False)
    observer.start()

    try:
        vtt_file = transcripts_dir / "test 01-01-2026.vtt"
        vtt_file.write_text("WEBVTT\n\n00:00:01.000 --> 00:00:02.000\nHello\n")
        triggered.wait(timeout=5)
    finally:
        observer.stop()
        observer.join()

    assert triggered.is_set()
    assert triggered_path[0] == vtt_file


def test_should_watcher_still_trigger_on_new_audio_file(tmp_path):
    import threading

    from watchdog.observers import Observer

    from call_intel.watcher import RecordingHandler

    recordings_dir = tmp_path / "recordings"
    recordings_dir.mkdir()

    triggered = threading.Event()
    triggered_path = []

    original_on_created = RecordingHandler.on_created

    class TestRecordingHandler(RecordingHandler):
        def on_created(self, event):
            path = Path(event.src_path)
            if path.suffix.lower() in {".m4a", ".wav", ".mp3", ".mp4", ".webm", ".ogg", ".flac"}:
                triggered_path.append(path)
                triggered.set()

    handler = TestRecordingHandler(skip_google=True)
    observer = Observer()
    observer.schedule(handler, str(recordings_dir), recursive=False)
    observer.start()

    try:
        audio_file = recordings_dir / "test 01-01-2026.m4a"
        audio_file.write_bytes(b"\x00" * 100)
        triggered.wait(timeout=5)
    finally:
        observer.stop()
        observer.join()

    assert triggered.is_set()
    assert triggered_path[0] == audio_file
